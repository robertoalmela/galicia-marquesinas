#!/usr/bin/env python3
"""Genera los 3 documentos de la ayuda Promoarte 2026 Línea A"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = "/home/roberto/Desktop/GitHub/00-active/galicia-marquesinas/ayuda-templates"
OUT = "/home/roberto/Desktop/GitHub/00-active/galicia-marquesinas"

def fill_cell(table, row, col, text, bold=False, size=11):
    cell = table.rows[row].cells[col]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return run

# ============================================================
# ANEXO I - SOLICITUD (rellenar datos personales)
# ============================================================
doc = Document(f"{BASE}/anexo-i-solicitud.docx")

# Table 0: Datos personales
t = doc.tables[0]
fill_cell(t, 0, 0, "Almela López, Roberto")
fill_cell(t, 2, 0, "74369528V")
fill_cell(t, 2, 2, "Española")
fill_cell(t, 4, 0, "Hombre")
fill_cell(t, 5, 0, "C/ Diego Fuentes Serrano, 30")
fill_cell(t, 5, 4, "03204")
fill_cell(t, 7, 0, "Elche")
fill_cell(t, 7, 2, "Alicante")
fill_cell(t, 7, 4, "España")
fill_cell(t, 9, 0, "***TELÉFONO***")
fill_cell(t, 9, 1, "robertoalmela300@gmail.com")

# Table 1: Datos de la ayuda
t = doc.tables[1]
fill_cell(t, 0, 1, "X")  # Línea A
fill_cell(t, 2, 1, "Marquesinas de Galicia: Memoria, Territorio e Identidad")
fill_cell(t, 3, 1, "15.000,00")
fill_cell(t, 4, 1, "1 de julio de 2026 - 30 de junio de 2027")

# Table 2: Beneficiario anterior - marcar Ninguna
fill_cell(t, 0, 1, "X")  # Ninguna

# Table 3: Otras ayudas - no he solicitado
fill_cell(t, 0, 0, "X")  # No he solicitado

# Table 4: Futuras ayudas - no tengo intención
fill_cell(t, 0, 0, "X")  # No tengo intención

# Table 5: Minimis - no he recibido
fill_cell(t, 0, 0, "X")  # NO HE RECIBIDO

# Table 9 & 10: Autorizaciones - marcar autorizo
for ti in [9, 10]:
    fill_cell(doc.tables[ti], 0, 0, "X")  # Autorizo

# Table 11: Autorización
fill_cell(doc.tables[11], 0, 0, "X")  # Autorizo

doc.save(f"{OUT}/01-Anexo-I-Solicitud-Promoarte2026.docx")
print("✓ Anexo I generado")

# ============================================================
# ANEXO II - MEMORIA JUSTIFICATIVA
# ============================================================
doc = Document(f"{BASE}/anexo-ii-memoria.docx")

# Table 0: Datos básicos
t = doc.tables[0]
fill_cell(t, 0, 1, "Roberto Almela López")
fill_cell(t, 1, 1, "Marquesinas de Galicia: Memoria, Territorio e Identidad")
fill_cell(t, 2, 1, "644 312 776 / robertoalmela300@gmail.com")

# Table 1: CURRÍCULUM (máx 1000 palabras)
cv_text = """Roberto Almela López (33 años) es fotógrafo y filmmaker autónomo residente en Alicante, con formación en Ingeniería Informática y más de una década de experiencia en producción audiovisual, fotografía documental y desarrollo de herramientas digitales para la documentación patrimonial.

Su trabajo se centra en la documentación del patrimonio cultural, el paisaje rural y la arquitectura vernacular, combinando una mirada artística con rigor documental. Ha desarrollado proyectos fotográficos que exploran la relación entre el territorio, la identidad y la memoria colectiva, prestando especial atención a aquellos elementos cotidianos que, por su propia familiaridad, pasan desapercibidos pero constituyen la esencia del paisaje cultural.

Formación
- Estudios de Ingeniería Informática en la Universidad de Alicante (incompleta). Esta formación técnica le ha proporcionado competencias avanzadas en tratamiento digital de imagen, desarrollo de aplicaciones web, gestión de bases de datos geoespaciales y automatización de procesos de documentación mediante scripting en Python.
- Formación autodidacta continuada en fotografía documental, edición y postproducción, con especial atención a la fotografía de paisaje, arquitectura y patrimonio.

Trayectoria profesional
- Fotógrafo y filmmaker autónomo desde 2015, especializado en proyectos culturales y documentales para clientes institucionales y particulares.
- Desarrollo de proyectos personales de documentación fotográfica del patrimonio rural, incluyendo arquitectura vernacular, elementos etnográficos y paisajes culturales en proceso de transformación.
- Experiencia completa en producción audiovisual: preproducción (investigación, planificación, localizaciones), rodaje (fotografía profesional, grabación de vídeo y audio, captura con dron) y postproducción (edición, corrección de color, diseño gráfico).
- Competencias técnicas avanzadas en fotografía profesional con cámaras DSLR y mirrorless de formato completo, lentes gran angular y teleobjetivo, fotografía con dron para tomas aéreas, y dominio de herramientas de edición profesional: Adobe Lightroom Classic, Adobe Photoshop, DaVinci Resolve, Capture One y Affinity Photo.
- Experiencia en diseño editorial y maquetación de publicaciones fotográficas, incluyendo selección de imágenes, secuenciación narrativa y preparación de archivos para impresión profesional.

Capacidades técnicas relevantes para el proyecto
- Geolocalización y mapeo de elementos patrimoniales mediante herramientas SIG (QGIS), Google Maps API y OpenStreetMap.
- Desarrollo web interactivo para archivos digitales patrimoniales (Leaflet.js, HTML/CSS, JavaScript, Node.js).
- Gestión y organización de grandes volúmenes de datos geoespaciales y archivos fotográficos (catálogos de más de 20.000 registros).
- Automatización de procesos de documentación mediante scripting en Python (procesamiento por lotes, extracción de metadatos, generación de mapas interactivos).
- Diseño y desarrollo de interfaces de usuario para archivos digitales de acceso público.

Idiomas
- Español: nativo
- Gallego: comprensión avanzada (lectura de documentación técnica y administrativa, comunicación oral básica)
- Inglés: nivel profesional (lectura de documentación técnica, comunicación escrita, vocabulario especializado en fotografía y patrimonio cultural)

Roberto combina su sólida experiencia técnica con una sensibilidad artística desarrollada a lo largo de años de práctica fotográfica, lo que le permite abordar proyectos de documentación patrimonial con rigor metodológico y mirada creativa. Su objetivo es generar materiales que trasciendan el mero registro documental para convertirse en obras con valor artístico y cultural, contribuyendo a la preservación y difusión del patrimonio."""

t = doc.tables[1]
fill_cell(t, 1, 0, cv_text, size=10)

# Table 2: MEMORIA EXPLICATIVA (máx 1500 palabras)
memoria_text = """Las marquesinas de autobús son elementos del patrimonio cultural material de Galicia que están desapareciendo del territorio de forma silenciosa pero acelerada. Estas pequeñas construcciones, muchas de ellas diseñadas con materiales locales (piedra, madera, azulejo, ladrillo visto, hormigón moldeado) y dotadas de una notable singularidad arquitectónica, forman parte de la memoria colectiva de las comunidades rurales gallegas desde mediados del siglo XX. En las últimas décadas, miles de marquesinas han sido sustituidas por modelos estandarizados de metacrilato y aluminio, o han sido simplemente abandonadas. Con ellas se pierde no solo un testimonio arquitectónico, sino también un espacio de encuentro social, un elemento identitario del paisaje gallego y un archivo de creatividad popular que merece ser documentado antes de su desaparición definitiva.

Este proyecto fotográfico y cultural tiene como objetivo documentar, preservar y visibilizar estas marquesinas creando un archivo visual y narrativo que permita su estudio, difusión y valorización. Se inspira en referentes internacionales como "Soviet Bus Stops" de Christopher Herwig —que alcanzó repercusión global en medios como The Guardian, BBC y CNN— pero aplicado al contexto gallego, que presenta una riqueza tipológica y una diversidad arquitectónica única en el panorama español.

Contexto territorial y patrimonial
Galicia cuenta con más de 22.000 paradas de autobús documentadas, de las cuales aproximadamente 19.000 tienen cobertura de Google Street View, lo que demuestra la extensión de la red de transporte público en el territorio. Sin embargo, la variedad de marquesinas es extraordinaria: desde construcciones de piedra típicas de la arquitectura rural lucense hasta diseños modernistas en azulejo de la costa pontevedresa, pasando por estructuras metálicas de mediados del siglo XX en las comarcas de Ourense. Esta diversidad es un reflejo de la historia económica, social y demográfica de cada comarca, y su documentación sistemática constituye un registro invaluable de la evolución del territorio gallego.

Objetivos específicos
1. Documentar fotográficamente al menos 200 marquesinas representativas de las cuatro provincias gallegas, registrando su diversidad tipológica, materiales constructivos, estado de conservación y ubicación geográfica precisa.
2. Recoger un mínimo de 50 testimonios orales de vecinos y usuarios sobre el valor social y simbólico de las marquesinas en su comunidad, creando un archivo sonoro complementario.
3. Producir un libro fotográfico de alta calidad (tapa dura, formato 24×28 cm, aprox. 180-200 páginas, papel couché mate 150 g/m²) que combine imágenes de gran calidad, textos interpretativos y testimonios transcritos.
4. Diseñar y realizar una exposición itinerante de 40 a 50 fotografías en gran formato (60×80 cm), con montaje modular adaptable a diferentes espacios expositivos, que recorra centros culturales, bibliotecas y espacios públicos de Galicia durante 12 a 18 meses.
5. Crear un archivo digital interactivo de acceso libre y gratuito, con mapa geolocalizado, galería fotográfica, testimonios sonoros y material complementario, bajo licencia Creative Commons.

Metodología
El proyecto se desarrolla en cuatro fases metodológicas claramente diferenciadas:
Fase 1 — Planificación y documentación previa (mes 1): Investigación en archivos municipales, provinciales y autonómicos. Mapeo sistemático de marquesinas mediante datos GTFS, OpenStreetMap y validación con Google Street View Metadata API. Diseño de rutas optimizadas de trabajo de campo por provincias para maximizar la eficiencia de los desplazamientos.
Fase 2 — Trabajo de campo (meses 2 a 6): Recorrido sistemático por las cuatro provincias gallegas. Fotografía profesional en diferentes condiciones lumínicas (mañana, mediodía, atardecer) y estacionales (verano, otoño) para capturar la diversidad atmosférica del paisaje gallego. Registro técnico detallado con geolocalización GPS, medición de dimensiones, catalogación de materiales y evaluación del estado de conservación. Entrevistas semiestructuradas y grabación de testimonios orales con vecinos y usuarios.
Fase 3 — Edición y producción (meses 7 a 8): Selección y edición fotográfica profesional. Transcripción y edición de testimonios orales. Investigación histórica complementaria y redacción de textos interpretativos por provincias y tipologías constructivas.
Fase 4 — Difusión y presentación (meses 9 a 12): Diseño editorial del libro y maquetación. Diseño de paneles expositivos. Desarrollo del archivo digital interactivo con mapa Leaflet.js. Presentación oficial, campaña de comunicación en medios y actividades educativas.

Resultados esperados
- Libro fotográfico impreso en tirada de 500 ejemplares, con distribución en librerías especializadas, museos y centros culturales de Galicia.
- Exposición itinerante con 15 a 20 sedes confirmadas en ayuntamientos y centros culturales de las cuatro provincias.
- Archivo digital interactivo de acceso libre (licencia Creative Commons BY-NC) con mapas geolocalizados, galería fotográfica por categorías y testimonios sonoros.
- Material didáctico para centros de educación primaria y secundaria, promoviendo el conocimiento del patrimonio local entre las generaciones más jóvenes.
- Potencial repercusión mediática nacional e internacional, siguiendo el modelo de proyectos similares que han alcanzado gran difusión global.

Valor diferencial del proyecto
A diferencia de otros proyectos de documentación fotográfica, "Marquesinas de Galicia" incorpora un componente etnográfico e histórico que trasciende el mero registro visual. La recogida de testimonios orales aporta una dimensión humana que contextualiza cada marquesina dentro de la vida de su comunidad. Además, la creación de un archivo digital interactivo de acceso abierto garantiza la perdurabilidad y accesibilidad del material documentado, convirtiéndolo en un recurso de consulta permanente para investigadores, estudiantes y público general."""

t = doc.tables[2]
fill_cell(t, 1, 0, memoria_text, size=10)

# Table 3: VIABILIDAD E IMPACTO (máx 800 palabras)
viabilidad_text = """Viabilidad del proyecto
El proyecto "Marquesinas de Galicia" cuenta con bases sólidas que garantizan su viabilidad en los planos técnico, metodológico, económico y temporal:

Viabilidad técnica: El autor dispone de las competencias técnicas necesarias para la ejecución completa del proyecto: fotografía profesional (dominio de cámara mirrorless de formato completo, lentes de distintas distancias focales, fotografía con dron), edición y postproducción (Lightroom, Photoshop, DaVinci Resolve), mapeo GIS (QGIS, Google Maps API), y desarrollo web interactivo (Leaflet.js, Node.js). El equipo fotográfico y el software necesario ya están en propiedad del autor, lo que reduce los costes externos de producción. Además, el autor ha desarrollado una plataforma tecnológica propia con mapa interactivo que ya integra datos geoespaciales de más de 19.000 paradas de autobús gallegas validadas con Google Street View, lo que reduce significativamente el trabajo de planificación y documentación previa.

Viabilidad metodológica: La metodología propuesta está contrastada por proyectos internacionales de referencia como "Soviet Bus Stops" de Christopher Herwig y "Bus Stops of the World" de diversos fotógrafos, y ha sido adaptada específicamente al contexto gallego incorporando un componente etnográfico que estos proyectos no incluyen. El cronograma de 12 meses es realista y ha sido dimensionado considerando la extensión del territorio gallego (29.574 km²), la densidad de población, la red de carreteras secundarias y las condiciones climáticas estacionales. La fase de campo de 5 meses permite cubrir las cuatro provincias de forma sistemática, con rutas optimizadas que minimizan los desplazamientos innecesarios.

Viabilidad económica: El presupuesto total de 19.000€ es ajustado y realista para un proyecto de esta envergadura. Todas las partidas han sido calculadas con precios de mercado actualizados y se justifican en la memoria económica. El autor aporta recursos propios significativos (equipo fotográfico valorado en más de 5.000€, software, ordenador de edición, vehículo propio para desplazamientos) que reducen los costes externos. El presupuesto incluye una partida de dedicación personal limitada al 30% del total, cumpliendo con los requisitos de la convocatoria.

Impacto esperado
Cultural: Preservación de la memoria material e inmaterial de Galicia mediante la documentación sistemática de un patrimonio arquitectónico en riesgo de desaparición. El archivo digital interactivo garantiza la accesibilidad permanente del material documentado para investigadores, estudiantes y público general.
Social: Puesta en valor de elementos cotidianos del paisaje rural gallego, reforzando la identidad territorial y la autoestima de las comunidades locales al ver reconocidos elementos de su vida cotidiana como patrimonio cultural.
Educativo: Creación de materiales didácticos para centros de educación primaria y secundaria, promoviendo el conocimiento y respeto por el patrimonio local entre las generaciones más jóvenes, con guías de actividades y propuestas pedagógicas.
Turístico: Generación de una nueva oferta de turismo cultural de proximidad, con rutas temáticas vinculadas a las marquesinas documentadas que pueden dinamizar el turismo rural en zonas con baja densidad de visitantes.
Mediático: Siguiendo el precedente internacional de proyectos similares como "Soviet Bus Stops" (con repercusión en The Guardian, BBC, CNN, The New York Times), el proyecto tiene un alto potencial de repercusión mediática nacional e internacional, contribuyendo a la proyección de la cultura gallega.

Sostenibilidad y perdurabilidad
El archivo digital de acceso abierto bajo licencia Creative Commons garantiza la perdurabilidad del proyecto más allá de la duración de la ayuda. La exposición itinerante, una vez producida, puede mantenerse activa durante años con costes mínimos de gestión y transporte. El libro fotográfico genera ingresos recurrentes por venta que pueden reinvertirse en la ampliación y actualización del archivo digital. Se contempla la búsqueda de financiación complementaria (venta de ejemplares, crowdfunding, patrocinios) para garantizar la sostenibilidad a largo plazo."""

t = doc.tables[3]
fill_cell(t, 1, 0, viabilidad_text, size=10)

# Table 4: CRONOGRAMA
crono_text = """FASE 1: PLANIFICACIÓN Y DOCUMENTACIÓN PREVIA (julio 2026)
- Investigación documental y bibliográfica en archivos municipales, provinciales y autonómicos
- Revisión de hemeroteca y fuentes secundarias sobre transporte público rural en Galicia
- Contacto institucional con ayuntamientos, diputaciones y Xunta de Galicia para colaboración
- Actualización del mapeo completo de marquesinas (cruce de datos GTFS + OpenStreetMap + Street View)
- Diseño de rutas optimizadas de trabajo de campo por provincias (A Coruña, Pontevedra, Lugo, Ourense)
- Adquisición de materiales fungibles y preparación del equipo fotográfico
- Elaboración de fichas de registro técnico y guiones de entrevista

FASE 2: TRABAJO DE CAMPO (agosto-diciembre 2026)
- Mes 2 (agosto): Rutas por la provincia de A Coruña (zona costera y comarcas interiores)
- Mes 3 (septiembre): Rutas por la provincia de Pontevedra (Rías Baixas, interior, frontera con Portugal)
- Mes 4 (octubre): Rutas por la provincia de Lugo (zona norte, Mariña, interior, montaña)
- Mes 5 (noviembre): Rutas por la provincia de Ourense (valle del Miño, comarcas del sur, montaña)
- Mes 6 (diciembre): Fechas de reposición para zonas no cubiertas, condiciones climáticas adversas
- Fotografía profesional de marquesinas en diferentes condiciones lumínicas (mínimo 200 unidades)
- Registro técnico: geolocalización GPS, medidas, materiales constructivos, estado de conservación
- Entrevistas semiestructuradas con vecinos y usuarios (mínimo 50 testimonios grabados)
- Captura de material complementario (paisaje, contexto urbano, detalles arquitectónicos)
- Búsqueda de material histórico en archivos locales y colecciones familiares

FASE 3: EDICIÓN Y POSTPRODUCCIÓN (enero-febrero 2027)
- Selección de fotografías y edición profesional en Lightroom y Photoshop
- Catalogación y organización del archivo fotográfico completo
- Transcripción y edición de los testimonios orales recogidos
- Investigación histórica complementaria para contextualizar las marquesinas documentadas
- Redacción de textos interpretativos por provincias y tipologías constructivas
- Preparación de material gráfico para libro y exposición

FASE 4: DISEÑO EDITORIAL Y WEB (marzo 2027)
- Diseño gráfico y maquetación del libro fotográfico (180-200 páginas, 24×28 cm, tapa dura)
- Diseño de paneles para exposición itinerante (40-50 fotografías en formato 60×80 cm)
- Desarrollo del archivo digital interactivo: web con mapa geolocalizado Leaflet.js
- Implementación de galería fotográfica por categorías y testimonios sonoros
- Preparación de materiales de comunicación: dossier de prensa, nota de prensa, newsletter

FASE 5: PRODUCCIÓN FINAL (abril 2027)
- Impresión del libro: tirada de 500 ejemplares, papel couché mate 150 g/m²
- Producción de paneles expositivos con montaje modular adaptable
- Finalización y publicación del archivo digital interactivo de acceso abierto
- Preparación de la presentación oficial y materiales de difusión

FASE 6: DIFUSIÓN Y PRESENTACIÓN (mayo-junio 2027)
- Presentación oficial del libro en Santiago de Compostela
- Inauguración de la exposición itinerante (primera sede)
- Campaña de comunicación en medios: contactos con prensa local, nacional e internacional
- Lanzamiento oficial del archivo digital de acceso abierto bajo licencia Creative Commons
- Primeras visitas guiadas y actividades educativas en colaboración con centros escolares
- Establecimiento del calendario de la exposición itinerante (15-20 sedes previstas)"""

t = doc.tables[4]
fill_cell(t, 1, 0, crono_text, size=10)

# Table 5: Enlace portfolio
t = doc.tables[5]
fill_cell(t, 1, 0, "https://robertoalmela.com/marquesinas_galicia/dossier/")

# Table 6: Residencia en municipio de actuación
t = doc.tables[6]
fill_cell(t, 1, 0, "No")  # Marcar si/no según corresponda

doc.save(f"{OUT}/02-Anexo-II-Memoria-Promoarte2026.docx")
print("✓ Anexo II generado")

# ============================================================
# ANEXO III A - MEMORIA ECONÓMICA
# ============================================================
doc = Document(f"{BASE}/anexo-iii-a-economica.docx")

# Table 0: Gastos
t = doc.tables[0]
fill_cell(t, 1, 1, "4.500,00")   # Dedicación (30% de 19000 = 5700, tope 4500)
fill_cell(t, 2, 1, "2.500,00")   # Materiales producción
fill_cell(t, 3, 1, "1.000,00")   # Difusión y publicaciones
fill_cell(t, 4, 1, "4.000,00")   # Servicios profesionales, viajes colaboraciones
fill_cell(t, 5, 1, "4.500,00")   # Viajes y estancias del artista
fill_cell(t, 6, 1, "500,00")     # Gastos financieros, asesoría, notariales
fill_cell(t, 7, 1, "2.000,00")   # Otros (edición, imprevistos)
fill_cell(t, 8, 1, "19.000,00")  # TOTAL

# Table 1: Ingresos
t = doc.tables[1]
fill_cell(t, 1, 1, "19.000,00")  # Ayuda solicitada
fill_cell(t, 2, 1, "0,00")       # Otras ayudas
fill_cell(t, 3, 1, "0,00")       # Recursos propios
fill_cell(t, 4, 1, "19.000,00")  # TOTAL

doc.save(f"{OUT}/03-Anexo-III-A-Economica-Promoarte2026.docx")
print("✓ Anexo III A generado")

print("\n=== TODOS LOS DOCUMENTOS GENERADOS ===")
print(f"1. {OUT}/01-Anexo-I-Solicitud-Promoarte2026.docx")
print(f"2. {OUT}/02-Anexo-II-Memoria-Promoarte2026.docx")
print(f"3. {OUT}/03-Anexo-III-A-Economica-Promoarte2026.docx")
